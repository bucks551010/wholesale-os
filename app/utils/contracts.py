import os, sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

import io
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Environment, BaseLoader

# ── Jinja2 env for text templates ─────────────────────────────────────────────
_jinja = Environment(loader=BaseLoader())

YELLOW_LETTER_TEMPLATE = """\
{{ today }}

Dear {{ owner_name | title }},

My name is {{ buyer_name }} and I'm a local real estate investor.
I am interested in purchasing your property at:

    {{ address }}
    {{ city }}, TX  {{ zip_code }}

I buy houses AS-IS, in any condition, with no agents, no fees,
and a flexible closing date that works for you.

If you have any interest in a quick, hassle-free sale, please
call or text me at {{ buyer_phone }}.

There is no obligation — I'd simply love to make you an offer.

Sincerely,

{{ buyer_name }}
{{ buyer_phone }}
{{ buyer_email }}

P.S. — Even if you're not ready to sell today, please keep my
number. I'm always looking and can close in as little as 10 days.
"""


def render_yellow_letter(owner_name: str, address: str, city: str,
                         zip_code: str, buyer_name: str,
                         buyer_phone: str, buyer_email: str) -> str:
    tmpl = _jinja.from_string(YELLOW_LETTER_TEMPLATE)
    return tmpl.render(
        today=date.today().strftime("%B %d, %Y"),
        owner_name=owner_name or "Property Owner",
        address=address, city=city, zip_code=zip_code,
        buyer_name=buyer_name, buyer_phone=buyer_phone,
        buyer_email=buyer_email,
    )


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_field(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(str(value or "___________________________"))


def generate_purchase_contract(deal: dict, parcel: dict, owner: dict,
                                buyer_name: str, buyer_entity: str) -> bytes:
    """
    Generate a simplified Texas purchase & sale agreement as DOCX bytes.
    ⚠ NOT LEGAL ADVICE — have a licensed attorney review before use.
    """
    doc = Document()
    doc.core_properties.title = "Real Estate Purchase Agreement"

    # Title
    _add_heading(doc, "REAL ESTATE PURCHASE AND SALE AGREEMENT")
    doc.add_paragraph(
        "⚠ THIS IS A TEMPLATE FOR EDUCATIONAL PURPOSES ONLY. "
        "CONSULT A LICENSED TEXAS REAL ESTATE ATTORNEY BEFORE USING.",
        style="Intense Quote",
    )
    doc.add_paragraph()

    today = date.today()
    closing = deal.get("closing_date") or (today + timedelta(days=30))
    option_days = deal.get("option_period_days") or 10
    option_exp = today + timedelta(days=option_days)
    purchase_price = deal.get("purchase_price") or 0
    em = deal.get("earnest_money_amount") or 500

    _add_field(doc, "Effective Date", today.strftime("%B %d, %Y"))
    doc.add_paragraph()

    _add_heading(doc, "PARTIES", level=2)
    _add_field(doc, "Seller", owner.get("owner_name") or deal.get("seller_name") or "")
    _add_field(doc, "Buyer", f"{buyer_name}  /  {buyer_entity}")
    doc.add_paragraph()

    _add_heading(doc, "PROPERTY", level=2)
    _add_field(doc, "Address", parcel.get("full_address") or "")
    _add_field(doc, "Legal Description", f"Parcel ID {parcel.get('parcel_id', '')}, Harris County, TX")
    doc.add_paragraph()

    _add_heading(doc, "PURCHASE PRICE & TERMS", level=2)
    _add_field(doc, "Purchase Price", f"${purchase_price:,.2f}")
    _add_field(doc, "Earnest Money", f"${em:,.2f} (due within 3 days of execution)")
    _add_field(doc, "Title Company", deal.get("title_company") or "")
    _add_field(doc, "Closing Date", str(closing))
    doc.add_paragraph()

    _add_heading(doc, "OPTION PERIOD", level=2)
    p = doc.add_paragraph()
    p.add_run(
        f"Buyer shall have an unrestricted right to terminate this contract "
        f"for any reason within {option_days} days of the effective date "
        f"(Option Period expires {option_exp.strftime('%B %d, %Y')}). "
        f"Option fee: $10.00, non-refundable, credited at closing."
    )
    doc.add_paragraph()

    _add_heading(doc, "CONDITION / AS-IS", level=2)
    doc.add_paragraph(
        "Buyer accepts the property in its current AS-IS condition. "
        "Seller makes no warranties regarding condition, habitability, "
        "or compliance with any laws or regulations."
    )
    doc.add_paragraph()

    _add_heading(doc, "SPECIAL PROVISIONS", level=2)
    doc.add_paragraph(deal.get("notes") or "(none)")
    doc.add_paragraph()

    _add_heading(doc, "SIGNATURES", level=2)
    for party in ["SELLER", "BUYER"]:
        doc.add_paragraph(f"{party}: ___________________________    Date: __________")
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_assignment_contract(deal: dict, parcel: dict,
                                 assignor: str, assignee: str,
                                 assignment_fee: float) -> bytes:
    """
    Generate a contract assignment agreement as DOCX bytes.
    ⚠ NOT LEGAL ADVICE — have a licensed attorney review before use.
    """
    doc = Document()
    _add_heading(doc, "CONTRACT ASSIGNMENT AGREEMENT")
    doc.add_paragraph(
        "⚠ TEMPLATE ONLY — CONSULT A LICENSED ATTORNEY BEFORE USE.",
        style="Intense Quote",
    )
    doc.add_paragraph()

    today = date.today()
    purchase_price = deal.get("purchase_price") or 0
    closing = deal.get("closing_date") or (today + timedelta(days=14))

    _add_field(doc, "Date", today.strftime("%B %d, %Y"))
    doc.add_paragraph()

    _add_heading(doc, "PARTIES", level=2)
    _add_field(doc, "Assignor (original buyer)", assignor)
    _add_field(doc, "Assignee (new buyer)", assignee)
    doc.add_paragraph()

    _add_heading(doc, "PROPERTY & ORIGINAL CONTRACT", level=2)
    _add_field(doc, "Property", parcel.get("full_address") or "")
    _add_field(doc, "Parcel ID", parcel.get("parcel_id") or "")
    _add_field(doc, "Original Purchase Price", f"${purchase_price:,.2f}")
    _add_field(doc, "Closing Date", str(closing))
    doc.add_paragraph()

    _add_heading(doc, "ASSIGNMENT FEE", level=2)
    _add_field(doc, "Assignment Fee", f"${assignment_fee:,.2f}")
    doc.add_paragraph(
        "Assignment fee is due at closing, paid by Assignee from closing proceeds. "
        "Assignee assumes all rights and obligations of the original purchase contract."
    )
    doc.add_paragraph()

    _add_heading(doc, "SIGNATURES", level=2)
    for party in ["ASSIGNOR", "ASSIGNEE"]:
        doc.add_paragraph(f"{party}: ___________________________    Date: __________")
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
